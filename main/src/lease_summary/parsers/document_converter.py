"""Document conversion utilities for non-PDF files.

Converts Word documents (.docx) to PDF for processing.
Uses python-docx for extraction and PyMuPDF for PDF generation.
"""
from __future__ import annotations

import tempfile
from pathlib import Path

try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# Supported file extensions for conversion
CONVERTIBLE_EXTENSIONS = {".docx"}


def is_convertible_document(filename: str | Path) -> bool:
    """Check if a file can be converted to PDF."""
    ext = Path(filename).suffix.lower()
    return ext in CONVERTIBLE_EXTENSIONS


def convert_to_pdf(input_path: str | Path, output_dir: str | Path | None = None) -> Path:
    """
    Convert a document to PDF.

    Args:
        input_path: Path to the document file (.docx, etc.)
        output_dir: Optional directory for output. If None, uses temp directory.

    Returns:
        Path to the generated PDF file.

    Raises:
        ValueError: If file type is not supported or required libraries missing.
        RuntimeError: If conversion fails.
    """
    input_path = Path(input_path)

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    ext = input_path.suffix.lower()

    if ext == ".docx":
        return _convert_docx_to_pdf(input_path, output_dir)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _convert_docx_to_pdf(docx_path: Path, output_dir: str | Path | None = None) -> Path:
    """Convert a Word document to PDF using python-docx + PyMuPDF."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx is required for Word document conversion")
    if not HAS_FITZ:
        raise RuntimeError("PyMuPDF is required for PDF generation")

    # Determine output path
    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = output_dir / f"{docx_path.stem}.pdf"
    else:
        pdf_path = Path(tempfile.gettempdir()) / f"{docx_path.stem}_{id(docx_path)}.pdf"

    try:
        # Load the Word document
        doc = Document(str(docx_path))

        # Create a new PDF
        pdf_doc = fitz.open()

        # Page settings
        page_width = 595  # A4 width in points
        page_height = 842  # A4 height in points
        margin = 72  # 1 inch margin
        line_height = 14  # Line height in points
        max_line_width = page_width - 2 * margin

        current_page = None
        current_y = margin

        def new_page():
            nonlocal current_page, current_y
            current_page = pdf_doc.new_page(width=page_width, height=page_height)
            current_y = margin

        def add_text(text: str, font_size: int = 11, bold: bool = False, italic: bool = False):
            nonlocal current_y

            if current_page is None:
                new_page()

            # Simple text wrapping
            font_name = "helv"  # Helvetica

            # Rough estimate of characters per line based on font size
            chars_per_line = int(max_line_width / (font_size * 0.6))

            words = text.split()
            lines = []
            current_line = []
            current_length = 0

            for word in words:
                word_len = len(word)
                if current_length + word_len + (1 if current_line else 0) <= chars_per_line:
                    current_line.append(word)
                    current_length += word_len + (1 if current_line else 0)
                else:
                    if current_line:
                        lines.append(" ".join(current_line))
                    current_line = [word]
                    current_length = word_len

            if current_line:
                lines.append(" ".join(current_line))

            for line in lines:
                if current_y + line_height > page_height - margin:
                    new_page()

                # Insert text
                text_point = fitz.Point(margin, current_y + font_size * 0.8)
                current_page.insert_text(
                    text_point,
                    line,
                    fontsize=font_size,
                    fontname=font_name,
                    color=(0, 0, 0)
                )
                current_y += line_height

        def add_paragraph_spacing():
            nonlocal current_y
            current_y += line_height * 0.5

        # Process document paragraphs
        for para in doc.paragraphs:
            if not para.text.strip():
                add_paragraph_spacing()
                continue

            # Determine style based on paragraph properties
            font_size = 11
            is_bold = False
            is_italic = False

            # Check paragraph style
            style_name = para.style.name.lower() if para.style else ""

            if "heading" in style_name or "title" in style_name:
                font_size = 16
                is_bold = True
            elif "heading 2" in style_name or "subtitle" in style_name:
                font_size = 14
                is_bold = True

            # Check runs for formatting
            full_text = ""
            for run in para.runs:
                text = run.text
                if run.bold:
                    is_bold = True
                if run.italic:
                    is_italic = True
                # Add superserscript/subscript handling if needed
                full_text += text

            if not full_text.strip():
                continue

            add_text(full_text.strip(), font_size=font_size, bold=is_bold, italic=is_italic)
            add_paragraph_spacing()

        # Process tables (basic support)
        for table in doc.tables:
            add_paragraph_spacing()

            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace("\n", " ")
                    row_texts.append(cell_text)

                if row_texts:
                    # Format as pipe-separated for simplicity
                    table_line = " | ".join(row_texts)
                    add_text(table_line, font_size=10)

            add_paragraph_spacing()

        # Save PDF
        pdf_doc.save(str(pdf_path))
        pdf_doc.close()

        return pdf_path

    except Exception as e:
        # Clean up partial file if it exists
        if pdf_path.exists():
            try:
                pdf_path.unlink()
            except:
                pass
        raise RuntimeError(f"Failed to convert Word document: {e}") from e


def get_converted_pdf_path(original_path: str | Path, temp_dir: str | Path | None = None) -> Path | None:
    """
    Get the path to a converted PDF, converting if necessary.

    Returns:
        Path to PDF (either original if already PDF, or converted).
        None if conversion not possible.
    """
    original_path = Path(original_path)

    # Already a PDF
    if original_path.suffix.lower() == ".pdf":
        return original_path

    # Try to convert
    if is_convertible_document(original_path):
        try:
            return convert_to_pdf(original_path, temp_dir)
        except Exception:
            return None

    return None
